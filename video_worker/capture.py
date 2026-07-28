from __future__ import annotations

import asyncio
import hashlib
import hmac
import os
import re
import shutil
import subprocess
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Callable
from urllib.parse import urljoin, urlparse


CAPTURE_WIDTH = 1920
CAPTURE_HEIGHT = 1080
CAPTURE_FPS = 30
CAPTURE_FILENAME = "capture.mp4"
DEFAULT_TARGET_URL = "https://rn-document-platform.onrender.com/"
DEFAULT_ALLOWED_HOST = "rn-document-platform.onrender.com"
DOCUMENT_PATH_RE = re.compile(r"^/documentos/\d+/$")
NORMALIZE_TIMEOUT_SECONDS = max(
    60, int(os.getenv("VIDEO_CAPTURE_NORMALIZE_TIMEOUT_SECONDS", "600"))
)
NORMALIZE_HEARTBEAT_SECONDS = max(
    5, int(os.getenv("VIDEO_CAPTURE_NORMALIZE_HEARTBEAT_SECONDS", "15"))
)

ProgressCallback = Callable[[str], None]


@dataclass(frozen=True)
class DemoIdentity:
    email: str
    password: str = field(repr=False)
    full_name: str = "Docente Exemplo"
    professional_name: str = "Docente Exemplo"


@dataclass(frozen=True)
class CaptureResult:
    duration_seconds: float
    capture_bytes: int
    capture_sha256: str
    docx_download_confirmed: bool

    def as_dict(self) -> dict[str, float | int | str | bool]:
        return asdict(self)


def validate_target_url(value: str | None) -> str:
    raw = (value or os.getenv("VIDEO_CAPTURE_TARGET_URL") or DEFAULT_TARGET_URL).strip()
    parsed = urlparse(raw)
    allowed_hosts = {
        host.strip().lower().rstrip(".")
        for host in os.getenv(
            "VIDEO_CAPTURE_ALLOWED_HOSTS", DEFAULT_ALLOWED_HOST
        ).split(",")
        if host.strip()
    }
    hostname = (parsed.hostname or "").lower().rstrip(".")
    if (
        parsed.scheme != "https"
        or not hostname
        or hostname not in allowed_hosts
        or parsed.username
        or parsed.password
        or (parsed.port not in (None, 443))
        or parsed.query
        or parsed.fragment
    ):
        raise ValueError("URL de captura não autorizada.")
    if parsed.path not in ("", "/"):
        raise ValueError("A captura deve começar na página inicial autorizada.")
    return f"https://{hostname}/"


def demo_identity(token: str, job_id: str) -> DemoIdentity:
    if not token:
        raise ValueError("Token do renderizador ausente.")
    digest = hmac.new(
        token.encode("utf-8"),
        f"rn-documentai-capture:{job_id}".encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    return DemoIdentity(
        email=f"video-demo-{digest[:18]}@example.com",
        password=f"Rn!{digest[18:42]}9a",
    )


def _run(command: list[str], *, timeout: int) -> None:
    try:
        completed = subprocess.run(
            command,
            check=False,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError(
            f"A normalização da captura excedeu o limite de {timeout} segundos."
        ) from exc
    if completed.returncode:
        stderr = completed.stderr[-3000:].strip()
        raise RuntimeError(f"Falha ao normalizar a captura: {stderr}")


def _media_duration(path: Path) -> float:
    completed = subprocess.run(
        [
            "ffprobe",
            "-v",
            "error",
            "-show_entries",
            "format=duration",
            "-of",
            "default=noprint_wrappers=1:nokey=1",
            str(path),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=60,
    )
    if completed.returncode:
        raise RuntimeError("Não foi possível validar a duração da captura.")
    try:
        duration = float(completed.stdout.strip())
    except ValueError as exc:
        raise RuntimeError("Duração inválida na captura.") from exc
    if duration <= 0:
        raise RuntimeError("A captura ficou vazia.")
    return duration


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        while chunk := source.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def _normalize_capture(source: Path, destination: Path) -> None:
    temporary = destination.with_suffix(".tmp.mp4")
    temporary.unlink(missing_ok=True)
    try:
        _run(
            [
                "ffmpeg",
                "-nostdin",
                "-hide_banner",
                "-loglevel",
                "error",
                "-filter_threads",
                "1",
                "-y",
                "-i",
                str(source),
                "-vf",
                (
                    f"scale={CAPTURE_WIDTH}:{CAPTURE_HEIGHT}:"
                    "force_original_aspect_ratio=decrease,"
                    f"pad={CAPTURE_WIDTH}:{CAPTURE_HEIGHT}:"
                    "(ow-iw)/2:(oh-ih)/2:color=0x071930,"
                    f"fps={CAPTURE_FPS},format=yuv420p"
                ),
                "-an",
                "-c:v",
                "libx264",
                "-preset",
                "ultrafast",
                "-crf",
                "20",
                "-threads",
                "1",
                "-r",
                str(CAPTURE_FPS),
                "-movflags",
                "+faststart",
                str(temporary),
            ],
            timeout=NORMALIZE_TIMEOUT_SECONDS,
        )
        temporary.replace(destination)
    except Exception:
        temporary.unlink(missing_ok=True)
        raise


async def _normalize_capture_with_heartbeat(
    source: Path,
    destination: Path,
    progress: ProgressCallback | None,
) -> None:
    task = asyncio.create_task(
        asyncio.to_thread(_normalize_capture, source, destination),
        name="normalize-capture",
    )
    while True:
        try:
            await asyncio.wait_for(
                asyncio.shield(task),
                timeout=NORMALIZE_HEARTBEAT_SECONDS,
            )
            return
        except TimeoutError:
            if progress:
                progress("normalizing_video")


def _create_synthetic_source(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("Fonte sintética — Imunologia", level=0)
    document.add_paragraph(
        "Material criado exclusivamente para a demonstração pública do RN DocumentAI. "
        "Não contém dados pessoais nem conteúdo institucional real."
    )
    document.add_heading("Ementa sintética", level=1)
    document.add_paragraph(
        "Fundamentos da resposta imune inata e adaptativa; células e órgãos do sistema "
        "imunológico; antígenos, anticorpos, imunidade celular, memória imunológica e "
        "aplicações clínicas introdutórias."
    )
    document.add_heading("Objetivos de aprendizagem", level=1)
    for item in (
        "Distinguir os componentes da imunidade inata e adaptativa.",
        "Relacionar mecanismos imunológicos a situações clínicas introdutórias.",
        "Interpretar esquemas básicos de resposta imune com rigor acadêmico.",
    ):
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("Orientações metodológicas", level=1)
    document.add_paragraph(
        "Combinar exposição dialogada, estudo de casos sintéticos, mapas conceituais "
        "e atividades de revisão formativa."
    )
    document.save(path)


async def _goto(page: Any, base_url: str, path: str) -> None:
    await page.goto(
        urljoin(base_url, path.lstrip("/")),
        wait_until="domcontentloaded",
        timeout=60_000,
    )
    await page.wait_for_timeout(900)


async def _submit(page: Any) -> None:
    await page.locator('button[type="submit"]').last.click()
    await page.wait_for_load_state("domcontentloaded")
    await page.wait_for_timeout(500)


async def _select_first_nonempty(page: Any, selector: str) -> None:
    options = page.locator(f"{selector} option")
    for index in range(await options.count()):
        value = await options.nth(index).get_attribute("value")
        if value:
            await page.locator(selector).select_option(value)
            return
    raise RuntimeError("Opção acadêmica esperada não foi encontrada.")


async def _login(page: Any, base_url: str, identity: DemoIdentity) -> bool:
    await _goto(page, base_url, "/entrar/")
    await page.locator('input[name="username"]').fill(identity.email)
    await page.locator('input[name="password"]').fill(identity.password)
    await _submit(page)
    return urlparse(page.url).path == "/painel/"


async def _login_or_register(
    page: Any, base_url: str, identity: DemoIdentity
) -> None:
    if await _login(page, base_url, identity):
        return

    await _goto(page, base_url, "/cadastro/")
    await page.locator('input[name="full_name"]').fill(identity.full_name)
    await page.locator('input[name="professional_name"]').fill(
        identity.professional_name
    )
    await page.locator('input[name="email"]').fill(identity.email)
    await page.locator('input[name="password1"]').fill(identity.password)
    await page.locator('input[name="password2"]').fill(identity.password)
    await page.locator('input[name="accept_terms"]').check()
    await _submit(page)
    if urlparse(page.url).path == "/painel/":
        return

    if await _login(page, base_url, identity):
        return
    raise RuntimeError("Não foi possível preparar a conta sintética de demonstração.")


async def _ensure_academic_data(page: Any, base_url: str, job_dir: Path) -> str | None:
    await _goto(page, base_url, "/academico/")
    body = await page.locator("body").inner_text()
    if "Universidade Exemplo" not in body:
        await _goto(page, base_url, "/academico/instituicao/")
        await page.locator('input[name="name"]').fill("Universidade Exemplo")
        await page.locator('input[name="acronym"]').fill("UEX")
        await page.locator('input[name="city"]').fill("Natal")
        await page.locator('input[name="state"]').fill("RN")
        await _submit(page)

    await _goto(page, base_url, "/academico/")
    body = await page.locator("body").inner_text()
    if "Medicina — Universidade Exemplo" not in body:
        await _goto(page, base_url, "/academico/curso/")
        await _select_first_nonempty(page, 'select[name="institution"]')
        await page.locator('input[name="name"]').fill("Medicina")
        await page.locator('input[name="level"]').fill("Graduação")
        await _submit(page)

    await _goto(page, base_url, "/academico/")
    body = await page.locator("body").inner_text()
    if "Imunologia — Medicina" not in body:
        await _goto(page, base_url, "/academico/disciplina/")
        await _select_first_nonempty(page, 'select[name="institution"]')
        await _select_first_nonempty(page, 'select[name="course"]')
        await page.locator('input[name="name"]').fill("Imunologia")
        await page.locator('input[name="workload"]').fill("60")
        await page.locator('input[name="semester"]').fill("2026.2")
        await page.locator('textarea[name="syllabus"]').fill(
            "Fundamentos da imunidade inata e adaptativa, células, órgãos linfoides, "
            "antígenos, anticorpos, imunidade celular e memória imunológica."
        )
        await page.locator('textarea[name="objectives"]').fill(
            "Compreender os mecanismos fundamentais da resposta imune e relacioná-los "
            "a situações clínicas introdutórias."
        )
        await page.locator('textarea[name="bibliography"]').fill(
            "Referências sintéticas preparadas exclusivamente para esta demonstração."
        )
        await _submit(page)

    source_title = "Fonte Sintética de Imunologia — Demonstração"
    await _goto(page, base_url, "/documentos/")
    body = await page.locator("body").inner_text()
    if source_title not in body:
        source_path = job_dir / "Fonte_Sintetica_Imunologia_Demonstracao.docx"
        _create_synthetic_source(source_path)
        await _goto(page, base_url, "/fontes/nova/")
        await _select_first_nonempty(page, 'select[name="institution"]')
        await _select_first_nonempty(page, 'select[name="discipline"]')
        await page.locator('input[name="title"]').fill(source_title)
        await page.locator('select[name="kind"]').select_option("OTHER")
        await page.locator('input[name="file"]').set_input_files(str(source_path))
        await _submit(page)

    await _goto(page, base_url, "/documentos/")
    links = page.locator('a[href^="/documentos/"]')
    for index in range(await links.count()):
        href = await links.nth(index).get_attribute("href")
        if href and DOCUMENT_PATH_RE.fullmatch(href):
            return href
    return None


async def _show_chapter(page: Any, title: str, hold_ms: int = 2600) -> None:
    await page.evaluate(
        """
        (title) => {
          document.getElementById("rn-capture-chapter")?.remove();
          const card = document.createElement("div");
          card.id = "rn-capture-chapter";
          card.setAttribute("role", "presentation");
          card.style.cssText = [
            "position:fixed",
            "z-index:2147483647",
            "top:84px",
            "left:42px",
            "max-width:760px",
            "padding:18px 24px",
            "border-left:8px solid #ffc107",
            "border-radius:8px",
            "background:rgba(7,25,48,.94)",
            "box-shadow:0 16px 42px rgba(0,0,0,.28)",
            "color:#fff",
            "font:700 30px/1.25 system-ui,sans-serif",
            "letter-spacing:-.02em"
          ].join(";");
          card.textContent = title;
          document.body.appendChild(card);
        }
        """,
        title,
    )
    await page.wait_for_timeout(hold_ms)
    await page.evaluate(
        'document.getElementById("rn-capture-chapter")?.remove()'
    )


async def _scroll(page: Any, ratio: float, hold_ms: int) -> None:
    await page.evaluate(
        "(ratio) => window.scrollTo({top: Math.max(0, "
        "(document.documentElement.scrollHeight - window.innerHeight) * ratio), "
        "behavior: 'smooth'})",
        ratio,
    )
    await page.wait_for_timeout(hold_ms)


async def _record_demo(
    page: Any,
    context: Any,
    base_url: str,
    auth_cookies: list[dict[str, Any]],
    existing_document_path: str | None,
    download_path: Path,
    generation_timeout_ms: int,
    progress: ProgressCallback | None = None,
) -> bool:
    def report(stage: str) -> None:
        if progress:
            progress(stage)

    report("recording_home")
    await _goto(page, base_url, "/")
    await _show_chapter(page, "RN DocumentAI: do contexto ao Plano de Ensino")
    await page.wait_for_timeout(4200)
    await _scroll(page, 0.26, 5000)
    await _scroll(page, 0.52, 5000)
    await _scroll(page, 0.78, 4200)

    report("recording_pricing")
    await _goto(page, base_url, "/precos/")
    await _show_chapter(page, "Planos transparentes e início gratuito")
    await _scroll(page, 0.38, 5500)
    await _scroll(page, 0.68, 4200)

    report("recording_signup")
    await _goto(page, base_url, "/cadastro/")
    await _show_chapter(page, "Cadastro seguro — sem exibir credenciais")
    await page.locator('input[name="full_name"]').focus()
    await page.wait_for_timeout(6500)
    await _scroll(page, 0.46, 3500)

    await context.add_cookies(auth_cookies)
    report("recording_dashboard")
    await _goto(page, base_url, "/painel/")
    await _show_chapter(page, "Conta de demonstração preparada com dados sintéticos")
    await page.wait_for_timeout(5000)
    await _scroll(page, 0.35, 5000)
    await _scroll(page, 0.72, 4500)

    report("recording_academics")
    await _goto(page, base_url, "/academico/")
    await _show_chapter(page, "Contexto acadêmico reutilizável")
    await page.wait_for_timeout(4200)
    await _scroll(page, 0.52, 5200)
    await _scroll(page, 0.88, 4200)

    report("recording_sources")
    await _goto(page, base_url, "/documentos/")
    await _show_chapter(page, "Fontes privadas em PDF ou DOCX")
    await _scroll(page, 0.72, 5500)

    report("recording_generation_form")
    await _goto(page, base_url, "/documentos/novo/")
    await _show_chapter(page, "Configuração da geração")
    await _select_first_nonempty(page, 'select[name="discipline"]')
    await page.locator('input[name="period"]').fill("2026.2")
    await page.locator('input[name="weeks"]').fill("20")
    await page.locator('textarea[name="methodology"]').fill(
        "Aulas dialogadas, estudos de caso sintéticos e mapas conceituais."
    )
    await page.locator('textarea[name="assessment"]').fill(
        "Avaliação formativa, atividade aplicada e síntese final."
    )
    source_checks = page.locator('input[name="sources"]')
    if await source_checks.count():
        await source_checks.first.check()
    await page.locator('textarea[name="notes"]').fill(
        "Priorizar coerência entre objetivos, conteúdo, metodologia e avaliação."
    )
    await page.wait_for_timeout(4200)
    await _scroll(page, 0.42, 5200)
    await _scroll(page, 0.82, 4200)

    report("generating_document")
    if existing_document_path:
        await _show_chapter(page, "Gerando a primeira versão estruturada", 4200)
        await _goto(page, base_url, existing_document_path)
    else:
        await page.locator("#generation-submit").click(no_wait_after=True)
        try:
            await page.locator("#generation-status").wait_for(
                state="visible", timeout=7000
            )
        except Exception:
            pass
        await page.wait_for_timeout(5500)
        try:
            await page.wait_for_url(
                re.compile(r"/documentos/\d+/$"),
                wait_until="domcontentloaded",
                timeout=generation_timeout_ms,
            )
        except Exception as exc:
            body = await page.locator("body").inner_text()
            if "Não foi possível gerar o documento" in body:
                raise RuntimeError(
                    "A geração do Plano de Ensino falhou no RN DocumentAI."
                ) from exc
            raise RuntimeError(
                "O RN DocumentAI não concluiu a geração dentro do prazo."
            ) from exc
        await page.wait_for_timeout(1200)

    report("recording_document")
    await _show_chapter(page, "Plano pronto para revisão docente")
    await page.wait_for_timeout(4500)
    await _scroll(page, 0.25, 5200)
    await _scroll(page, 0.52, 5200)
    await _scroll(page, 0.78, 5200)
    await _scroll(page, 1.0, 4800)

    report("downloading_docx")
    download_link = page.get_by_role("link", name=re.compile(r"Baixar .*\.docx"))
    if not await download_link.count():
        download_link = page.get_by_role("link", name="Baixar DOCX")
    async with page.expect_download(timeout=60_000) as download_info:
        await download_link.last.click()
    download = await download_info.value
    await download.save_as(str(download_path))
    await _show_chapter(page, "DOCX gerado e download confirmado")
    await page.wait_for_timeout(5500)
    return download_path.exists() and download_path.stat().st_size > 0


async def capture_video(
    *,
    job_dir: Path,
    target_url: str,
    token: str,
    job_id: str,
    progress: ProgressCallback | None = None,
) -> CaptureResult:
    from playwright.async_api import async_playwright

    base_url = validate_target_url(target_url)
    identity = demo_identity(token, job_id)
    destination = job_dir / CAPTURE_FILENAME
    if destination.exists():
        return CaptureResult(
            duration_seconds=round(_media_duration(destination), 3),
            capture_bytes=destination.stat().st_size,
            capture_sha256=_sha256(destination),
            docx_download_confirmed=(job_dir / "Plano_Ensino_Demonstracao.docx").exists(),
        )

    job_dir.mkdir(parents=True, exist_ok=True)
    recording_dir = job_dir / "capture-recording"
    shutil.rmtree(recording_dir, ignore_errors=True)
    recording_dir.mkdir(parents=True, exist_ok=False)
    download_path = job_dir / "Plano_Ensino_Demonstracao.docx"
    generation_timeout_ms = max(
        60_000, int(os.getenv("VIDEO_CAPTURE_GENERATION_TIMEOUT_SECONDS", "360")) * 1000
    )

    def report(stage: str) -> None:
        if progress:
            progress(stage)

    raw_video: Path | None = None
    report("launching_browser")
    async with async_playwright() as playwright:
        browser = await playwright.chromium.launch(
            headless=True,
            args=[
                "--no-sandbox",
                "--disable-dev-shm-usage",
                "--disable-background-networking",
                "--disable-component-update",
                "--disable-default-apps",
                "--disable-features=Translate,MediaRouter",
                "--disable-sync",
                "--metrics-recording-only",
                "--no-first-run",
            ],
        )
        try:
            bootstrap = await browser.new_context(
                viewport={"width": CAPTURE_WIDTH, "height": CAPTURE_HEIGHT},
                locale="pt-BR",
                timezone_id="America/Fortaleza",
                accept_downloads=True,
            )
            try:
                bootstrap_page = await bootstrap.new_page()
                report("authenticating_demo")
                await _login_or_register(bootstrap_page, base_url, identity)
                report("preparing_academic_data")
                existing_document_path = await _ensure_academic_data(
                    bootstrap_page, base_url, job_dir
                )
                auth_cookies = await bootstrap.cookies()
            finally:
                await bootstrap.close()

            report("recording_demo")
            recorded = await browser.new_context(
                viewport={"width": CAPTURE_WIDTH, "height": CAPTURE_HEIGHT},
                screen={"width": CAPTURE_WIDTH, "height": CAPTURE_HEIGHT},
                device_scale_factor=1,
                locale="pt-BR",
                timezone_id="America/Fortaleza",
                color_scheme="light",
                reduced_motion="no-preference",
                accept_downloads=True,
                record_video_dir=str(recording_dir),
                record_video_size={
                    "width": CAPTURE_WIDTH,
                    "height": CAPTURE_HEIGHT,
                },
            )
            video = None
            try:
                page = await recorded.new_page()
                video = page.video
                download_confirmed = await _record_demo(
                    page,
                    recorded,
                    base_url,
                    auth_cookies,
                    existing_document_path,
                    download_path,
                    generation_timeout_ms,
                    report,
                )
                report("closing_recording")
            finally:
                await recorded.close()
            if video is None:
                raise RuntimeError("O navegador não iniciou a gravação.")
            raw_video = Path(await video.path())
        finally:
            await browser.close()

    if raw_video is None or not raw_video.exists():
        raise RuntimeError("O navegador não produziu o arquivo de captura.")

    report("normalizing_video")
    await _normalize_capture_with_heartbeat(raw_video, destination, progress)
    report("validating_video")
    try:
        duration = _media_duration(destination)
    except Exception:
        destination.unlink(missing_ok=True)
        raise
    if duration < 90:
        destination.unlink(missing_ok=True)
        raise RuntimeError("A captura ficou curta demais para a produção final.")
    shutil.rmtree(recording_dir, ignore_errors=True)
    report("ready")
    return CaptureResult(
        duration_seconds=round(duration, 3),
        capture_bytes=destination.stat().st_size,
        capture_sha256=_sha256(destination),
        docx_download_confirmed=download_confirmed,
    )
