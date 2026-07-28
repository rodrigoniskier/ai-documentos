import io
import re
from unittest import skipUnless

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from docx import Document
from pypdf import PdfReader

from .document_conversion import convert_docx_bytes_to_pdf, libreoffice_available
from .document_export import build_project_docx, build_project_pdf
from .document_models import DocumentProject, DocumentTemplate
from .document_quality import audit_docx_structure
from .models import User
from .services import ensure_plans, provision_free_account
from .template_renderer import normalize_label


def _set_bullets(cell, values):
    cell.text = ""
    for value in values:
        cell.add_paragraph(value, style="List Bullet")


def institutional_template_upload():
    """Cria uma miniatura estrutural do modelo institucional analisado."""

    output = io.BytesIO()
    document = Document()
    document.add_heading("PLANO DE ENSINO", level=1)
    document.add_paragraph("2026.2")

    metadata = document.add_table(rows=4, cols=4)
    metadata.style = "Table Grid"
    metadata.cell(0, 0).merge(metadata.cell(0, 3)).text = "PLANO DE ENSINO"
    metadata.cell(1, 0).merge(metadata.cell(1, 1)).text = "Curso: Medicina"
    metadata.cell(1, 2).merge(metadata.cell(1, 3)).text = (
        "Componente: Mecanismos de Agressão Patológicos e de Defesa 2"
    )
    metadata.cell(2, 0).text = "Período: P4"
    metadata.cell(2, 1).text = "Turno: Vespertino"
    metadata.cell(2, 2).text = "Modalidade: Presencial"
    metadata.cell(2, 3).text = "Semestre: 2026.2"
    metadata.cell(3, 0).merge(metadata.cell(3, 1)).text = (
        "Professor responsável: Rodrigo Niskier Ferreira Barbosa"
    )
    metadata.cell(3, 2).text = "C/H semanal: 3h/a"
    metadata.cell(3, 3).text = "C/H semestral: 60h/a"

    ementa = document.add_table(rows=2, cols=1)
    ementa.style = "Table Grid"
    ementa.cell(0, 0).text = "EMENTA"
    ementa.cell(1, 0).text = "Ementa institucional anterior."

    objectives = document.add_table(rows=4, cols=2)
    objectives.style = "Table Grid"
    objectives.cell(0, 0).merge(objectives.cell(0, 1)).text = (
        "OBJETIVOS / COMPETÊNCIAS"
    )
    for row_index, label in enumerate(
        ("Cognitivos", "Habilidades", "Atitudes"), start=1
    ):
        objectives.cell(row_index, 0).text = label
        _set_bullets(objectives.cell(row_index, 1), [f"Item antigo de {label}."])

    content = document.add_table(rows=6, cols=3)
    content.style = "Table Grid"
    content.cell(0, 0).merge(content.cell(0, 2)).text = "CONTEÚDO"
    content.cell(1, 0).text = "UNIDADE"
    content.cell(1, 1).text = "C/H"
    content.cell(1, 2).text = "TÓPICOS"
    for index in range(4):
        row = index + 2
        numeral = ("I", "II", "III", "IV")[index]
        content.cell(row, 0).text = f"Unidade {numeral}: Unidade antiga"
        content.cell(row, 1).text = "15h/a"
        _set_bullets(content.cell(row, 2), ["Tópico antigo 1", "Tópico antigo 2"])

    for heading in ("ESTRATÉGIAS DE ENSINO", "RECURSOS DISPONÍVEIS", "AVALIAÇÃO"):
        table = document.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        table.cell(0, 0).text = heading
        _set_bullets(table.cell(1, 0), ["Item institucional antigo."])

    bibliography = document.add_table(rows=2, cols=2)
    bibliography.style = "Table Grid"
    bibliography.cell(0, 0).text = "BIBLIOGRAFIA BÁSICA"
    bibliography.cell(0, 1).text = "BIBLIOGRAFIA COMPLEMENTAR"
    _set_bullets(bibliography.cell(1, 0), ["Referência básica antiga."])
    _set_bullets(bibliography.cell(1, 1), ["Referência complementar antiga."])

    review = document.add_table(rows=1, cols=1)
    review.style = "Table Grid"
    review.cell(0, 0).text = "PONTOS PARA REVISÃO DOCENTE"
    review.cell(0, 0).add_paragraph("Ponto antigo.", style="List Bullet")

    document.sections[0].header.paragraphs[0].text = (
        "Pró-Reitoria Acadêmica / Coordenação do Curso"
    )
    document.sections[0].footer.paragraphs[0].text = (
        "Documento sujeito à revisão e aprovação docente/institucional."
    )
    document.save(output)
    return SimpleUploadedFile(
        "modelo-institucional.docx",
        output.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )


def project_content():
    return {
        "title": "Plano de POMC",
        "warnings": [],
        "resolved_fields": [
            {"label": "Curso", "value": "Medicina"},
            {"label": "Componente", "value": "POMC"},
            {"label": "Período", "value": "P4"},
            {"label": "Turno", "value": "Vespertino"},
            {"label": "Modalidade", "value": "Presencial"},
            {"label": "Semestre", "value": "2026.2"},
            {
                "label": "Professor responsável",
                "value": "Rodrigo Niskier Ferreira Barbosa",
            },
            {"label": "C/H semanal", "value": "3h/a"},
            {"label": "C/H semestral", "value": "60h/a"},
        ],
        "sections": [
            {
                "heading": "EMENTA",
                "body": "Ementa personalizada sem reconstrução da tabela.",
            },
            {
                "heading": "OBJETIVOS / COMPETÊNCIAS",
                "body": (
                    "Cognitivos\n- Compreender fundamentos integrados.\n"
                    "- Relacionar mecanismos e clínica.\n"
                    "Habilidades\n- Interpretar evidências diagnósticas.\n"
                    "Atitudes\n- Atuar de forma ética e colaborativa."
                ),
            },
            {
                "heading": "CONTEÚDO",
                "body": (
                    "UNIDADE I: Microbiologia Clínica | C/H: 20h/a\n"
                    "- Bactérias de interesse médico\n- Virologia clínica\n"
                    "UNIDADE II: Parasitologia Médica | C/H: 15h/a\n"
                    "- Parasitoses clínicas\n- Diagnóstico microscópico\n"
                    "UNIDADE III: Imunologia Clínica | C/H: 15h/a\n"
                    "- Imunopatologias\n- Casos clínicos\n"
                    "UNIDADE IV: Oncobiologia | C/H: 10h/a\n"
                    "- Imunoterapia\n- Casos em oncologia"
                ),
            },
            {
                "heading": "ESTRATÉGIAS DE ENSINO",
                "body": "- Aulas dialogadas.\n- Discussão de casos clínicos.",
            },
            {
                "heading": "RECURSOS DISPONÍVEIS",
                "body": "- Laboratório.\n- Ambiente Virtual de Aprendizagem.",
            },
            {
                "heading": "AVALIAÇÃO",
                "body": "- Avaliação teórica.\n- Avaliação prática.",
            },
            {
                "heading": "BIBLIOGRAFIA BÁSICA",
                "body": "- Referência básica A.\n- Referência básica B.",
            },
            {
                "heading": "BIBLIOGRAFIA COMPLEMENTAR",
                "body": "- Referência complementar A.\n- Referência complementar B.",
            },
            {
                "heading": "PONTOS PARA REVISÃO DOCENTE",
                "body": "- Confirmar aderência ao PPC.\n- Realizar revisão humana final.",
            },
        ],
    }


def normalized_pdf_text(payload):
    reader = PdfReader(io.BytesIO(bytes(payload)))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return re.sub(r"\s+", " ", text).strip(), len(reader.pages)


@override_settings(
    DOCUMENT_VISUAL_QA_ENABLED=False,
    DOCUMENT_VISUAL_QA_BLOCK_ON_FAILURE=True,
)
class StrictTemplatePipelineTests(TestCase):
    def setUp(self):
        ensure_plans()
        self.user = User.objects.create_user(
            email="strict@example.com",
            password="SenhaForte123!",
            full_name="Professora Estrita",
            professional_name="Profa. Estrita",
        )
        provision_free_account(self.user)
        self.template = DocumentTemplate.objects.create(
            owner=self.user,
            title="Modelo institucional",
            document_type="PLANO_ENSINO",
            file=institutional_template_upload(),
        )
        self.project = DocumentProject.objects.create(
            owner=self.user,
            template=self.template,
            document_type="PLANO_ENSINO",
            title="Plano de POMC",
            course_context="Medicina",
            institution_context="Curso de Medicina, semestre 2026.2.",
            # Estes valores conflitantes não podem prevalecer sobre resolved_fields.
            field_values={"periodo": "P1", "carga_horaria": "120h/a"},
            content=project_content(),
            status="ready",
        )

    def test_normalization_does_not_mutilate_regular_words(self):
        self.assertEqual(normalize_label("CONTEÚDO"), "conteudo")
        self.assertEqual(normalize_label("Cognitivos"), "cognitivos")
        self.assertEqual(normalize_label("IV - Avaliação"), "avaliacao")

    def test_renderer_preserves_geometry_and_places_content_once(self):
        self.template.file.open("rb")
        try:
            template_bytes = self.template.file.read()
        finally:
            self.template.file.close()

        generated = build_project_docx(self.project, watermark=False)
        generated.seek(0)
        generated_bytes = generated.read()
        report = audit_docx_structure(template_bytes, generated_bytes)
        self.assertTrue(report.passed, report.issues)

        document = Document(io.BytesIO(generated_bytes))
        self.assertEqual(len(document.tables), 9)
        self.assertEqual([len(table.rows) for table in document.tables], [4, 2, 4, 6, 2, 2, 2, 2, 1])

        metadata = document.tables[0]
        self.assertEqual(metadata.cell(1, 0).text, "Curso: Medicina")
        self.assertEqual(metadata.cell(1, 2).text, "Componente: POMC")
        self.assertEqual(metadata.cell(2, 0).text, "Período: P4")
        self.assertEqual(metadata.cell(3, 3).text, "C/H semestral: 60h/a")
        self.assertNotIn("P1", " ".join(cell.text for row in metadata.rows for cell in row.cells))
        self.assertNotIn("120h/a", " ".join(cell.text for row in metadata.rows for cell in row.cells))

        objectives = document.tables[2]
        self.assertEqual(objectives.cell(1, 0).text, "Cognitivos")
        self.assertNotIn("Compreender fundamentos", objectives.cell(1, 0).text)
        self.assertIn("Compreender fundamentos integrados.", objectives.cell(1, 1).text)
        self.assertIn("Interpretar evidências diagnósticas.", objectives.cell(2, 1).text)
        self.assertIn("Atuar de forma ética", objectives.cell(3, 1).text)

        content = document.tables[3]
        self.assertEqual(content.cell(1, 0).text, "UNIDADE")
        self.assertEqual(content.cell(1, 1).text, "C/H")
        self.assertEqual(content.cell(1, 2).text, "TÓPICOS")
        self.assertIn("Microbiologia Clínica", content.cell(2, 0).text)
        self.assertEqual(content.cell(2, 1).text, "20h/a")
        self.assertIn("Bactérias de interesse médico", content.cell(2, 2).text)
        content_text = "\n".join(
            cell.text
            for row in content.rows
            for cell in {id(item._tc): item for item in row.cells}.values()
        )
        self.assertEqual(content_text.upper().count("UNIDADE I:"), 1)
        self.assertEqual(content_text.upper().count("UNIDADE II:"), 1)
        self.assertNotIn("UNIDADE I:", content.cell(2, 2).text.upper())

        bibliography = document.tables[7]
        self.assertEqual(bibliography.cell(0, 0).text, "BIBLIOGRAFIA BÁSICA")
        self.assertEqual(
            bibliography.cell(0, 1).text, "BIBLIOGRAFIA COMPLEMENTAR"
        )
        self.assertIn("Referência básica A.", bibliography.cell(1, 0).text)
        self.assertIn("Referência complementar A.", bibliography.cell(1, 1).text)

        self.assertIn(
            "Pró-Reitoria Acadêmica",
            document.sections[0].header.paragraphs[0].text,
        )
        self.assertIn(
            "Documento sujeito à revisão",
            document.sections[0].footer.paragraphs[0].text,
        )

    @skipUnless(libreoffice_available(), "LibreOffice não disponível")
    def test_pdf_is_the_conversion_of_the_final_docx(self):
        generated_docx = build_project_docx(self.project, watermark=False)
        generated_docx.seek(0)
        docx_bytes = generated_docx.read()

        expected_pdf = convert_docx_bytes_to_pdf(docx_bytes)
        generated_pdf = build_project_pdf(self.project, watermark=False)
        generated_pdf.seek(0)
        pdf_bytes = generated_pdf.read()

        self.assertTrue(pdf_bytes.startswith(b"%PDF"))
        expected_text, expected_pages = normalized_pdf_text(expected_pdf)
        actual_text, actual_pages = normalized_pdf_text(pdf_bytes)
        self.assertEqual(actual_pages, expected_pages)
        self.assertEqual(actual_text, expected_text)
