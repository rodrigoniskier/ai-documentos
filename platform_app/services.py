import io
import json
import secrets
from datetime import datetime

from django.conf import settings
from django.core.files.base import ContentFile
from django.db import transaction
from django.db.models import Sum
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openai import OpenAI
from pypdf import PdfReader

from .models import CreditEntry, Plan, Subscription, Wallet


PLAN_DEFAULTS = {
    "FREE": {
        "name": "Gratuito",
        "price_label": "R$ 0",
        "initial_credits": 5,
        "monthly_credits": 0,
        "institution_limit": 1,
        "discipline_limit": 1,
        "source_limit": 3,
        "daily_limit": 4,
        "watermark": True,
        "display_order": 1,
        "active": True,
    },
    "PRO": {
        "name": "Pro",
        "price_label": "R$ 19,90/mês",
        "initial_credits": 0,
        "monthly_credits": 40,
        "institution_limit": 2,
        "discipline_limit": 10,
        "source_limit": 30,
        "daily_limit": 12,
        "watermark": False,
        "display_order": 2,
        "active": True,
    },
    "PREMIUM": {
        "name": "Ultra",
        "price_label": "R$ 49,90/mês",
        "initial_credits": 0,
        "monthly_credits": 120,
        "institution_limit": 5,
        "discipline_limit": 30,
        "source_limit": 100,
        "daily_limit": 30,
        "watermark": False,
        "display_order": 3,
        "active": True,
    },
}


def ensure_plans():
    for code, defaults in PLAN_DEFAULTS.items():
        Plan.objects.update_or_create(code=code, defaults=defaults)


def provision_free_account(user):
    ensure_plans()
    plan = Plan.objects.get(code="FREE")
    Subscription.objects.get_or_create(user=user, defaults={"plan": plan})
    Wallet.objects.get_or_create(user=user)
    grant_credits(
        user=user,
        amount=plan.initial_credits,
        idempotency_key=f"free-signup:{user.pk}",
        reason="Créditos iniciais do plano gratuito",
    )


@transaction.atomic
def grant_credits(user, amount: int, idempotency_key: str, reason: str):
    wallet, _ = Wallet.objects.select_for_update().get_or_create(user=user)
    if CreditEntry.objects.filter(idempotency_key=idempotency_key).exists():
        return wallet

    before = wallet.balance
    wallet.balance += amount
    wallet.save(update_fields=["balance", "updated_at"])
    CreditEntry.objects.create(
        user=user,
        wallet=wallet,
        kind="grant",
        amount=amount,
        balance_before=before,
        balance_after=wallet.balance,
        idempotency_key=idempotency_key,
        reason=reason,
    )
    return wallet


@transaction.atomic
def reserve_credits(user, amount: int, idempotency_key: str, reference: str = ""):
    wallet = Wallet.objects.select_for_update().get(user=user)
    if CreditEntry.objects.filter(idempotency_key=idempotency_key).exists():
        return wallet

    used_today = -(
        CreditEntry.objects.filter(
            user=user,
            kind="reserve",
            created_at__date=timezone.localdate(),
        ).aggregate(total=Sum("amount"))["total"]
        or 0
    )
    if used_today + amount > user.subscription.plan.daily_limit:
        raise ValueError("O limite diário de créditos foi atingido.")
    if wallet.balance < amount:
        raise ValueError("Saldo de créditos insuficiente.")

    before = wallet.balance
    wallet.balance -= amount
    wallet.save(update_fields=["balance", "updated_at"])
    CreditEntry.objects.create(
        user=user,
        wallet=wallet,
        kind="reserve",
        amount=-amount,
        balance_before=before,
        balance_after=wallet.balance,
        idempotency_key=idempotency_key,
        reason="Geração de plano de ensino",
        reference=reference,
    )
    return wallet


@transaction.atomic
def refund_credits(user, amount: int, idempotency_key: str, reference: str = ""):
    wallet = Wallet.objects.select_for_update().get(user=user)
    if CreditEntry.objects.filter(idempotency_key=idempotency_key).exists():
        return wallet

    before = wallet.balance
    wallet.balance += amount
    wallet.save(update_fields=["balance", "updated_at"])
    CreditEntry.objects.create(
        user=user,
        wallet=wallet,
        kind="refund",
        amount=amount,
        balance_before=before,
        balance_after=wallet.balance,
        idempotency_key=idempotency_key,
        reason="Devolução por erro técnico",
        reference=reference,
    )
    return wallet


def extract_source_text(source):
    try:
        source.file.open("rb")
        if source.file.name.lower().endswith(".pdf"):
            text = "\n".join(
                (page.extract_text() or "") for page in PdfReader(source.file).pages
            )
        else:
            from docx import Document as DocxDocument

            document = DocxDocument(source.file)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        source.extracted_text = text[:80000]
        source.status = "done"
    except Exception:
        source.extracted_text = ""
        source.status = "failed"
    source.save(update_fields=["extracted_text", "status"])


PLAN_SCHEMA = {
    "type": "object",
    "properties": {
        "ementa": {"type": "string"},
        "objetivo_geral": {"type": "string"},
        "objetivos_especificos": {"type": "array", "items": {"type": "string"}},
        "competencias": {"type": "array", "items": {"type": "string"}},
        "conteudo": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "unidade": {"type": "string"},
                    "topicos": {"type": "array", "items": {"type": "string"}},
                    "carga": {"type": "string"},
                },
                "required": ["unidade", "topicos", "carga"],
                "additionalProperties": False,
            },
        },
        "metodologia": {"type": "array", "items": {"type": "string"}},
        "avaliacao": {"type": "array", "items": {"type": "string"}},
        "recursos": {"type": "array", "items": {"type": "string"}},
        "bibliografia_basica": {"type": "array", "items": {"type": "string"}},
        "bibliografia_complementar": {
            "type": "array",
            "items": {"type": "string"},
        },
        "observacoes": {"type": "string"},
    },
    "required": [
        "ementa",
        "objetivo_geral",
        "objetivos_especificos",
        "competencias",
        "conteudo",
        "metodologia",
        "avaliacao",
        "recursos",
        "bibliografia_basica",
        "bibliografia_complementar",
        "observacoes",
    ],
    "additionalProperties": False,
}


def generate_plan_with_openai(payload):
    if not settings.OPENAI_API_KEY:
        raise RuntimeError("A OpenAI ainda não foi configurada.")

    client = OpenAI(api_key=settings.OPENAI_API_KEY, timeout=90)
    response = client.responses.create(
        model=settings.OPENAI_MODEL,
        store=False,
        input=[
            {
                "role": "system",
                "content": (
                    "Crie um plano de ensino superior coerente, claro e revisável pelo "
                    "professor. Use somente o contexto fornecido e não invente normas "
                    "institucionais ou referências inexistentes."
                ),
            },
            {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
        ],
        text={
            "format": {
                "type": "json_schema",
                "name": "plano_de_ensino",
                "strict": True,
                "schema": PLAN_SCHEMA,
            }
        },
    )
    usage = getattr(response, "usage", None)
    return (
        json.loads(response.output_text),
        getattr(usage, "input_tokens", 0) or 0,
        getattr(usage, "output_tokens", 0) or 0,
    )


def new_document_code():
    return f"RND-{datetime.now():%Y%m%d}-{secrets.token_hex(3).upper()}"


def build_plan_docx(document_record, discipline, data, watermark: bool):
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(11)

    institution = discipline.institution
    if institution.logo:
        try:
            institution.logo.open("rb")
            document.add_picture(institution.logo, width=Cm(3))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(institution.name.upper()).bold = True
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("PLANO DE ENSINO").bold = True

    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    details = [
        ("Curso", discipline.course.name),
        ("Disciplina", discipline.name),
        ("Carga horária", f"{discipline.workload} horas"),
        ("Professor", document_record.owner.professional_name),
    ]
    for row, (label, value) in zip(table.rows, details):
        row.cells[0].text = label
        row.cells[1].text = value

    sections = [
        ("1. Ementa", "ementa"),
        ("2. Objetivo geral", "objetivo_geral"),
        ("3. Objetivos específicos", "objetivos_especificos"),
        ("4. Competências e habilidades", "competencias"),
        ("5. Metodologia", "metodologia"),
        ("6. Avaliação", "avaliacao"),
        ("7. Recursos", "recursos"),
        ("8. Bibliografia básica", "bibliografia_basica"),
        ("9. Bibliografia complementar", "bibliografia_complementar"),
        ("10. Observações", "observacoes"),
    ]
    for heading_text, key in sections:
        document.add_heading(heading_text, level=2)
        value = data[key]
        if isinstance(value, list):
            for item in value:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph(value or "—")

    document.add_heading("Conteúdo programático", level=2)
    for unit in data["conteudo"]:
        paragraph = document.add_paragraph()
        paragraph.add_run(f'{unit["unidade"]} — {unit["carga"]}').bold = True
        for topic in unit["topicos"]:
            document.add_paragraph(topic, style="List Bullet")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = f"Documento {document_record.code}"
    if watermark:
        footer.add_run(" • Gerado com RN DocumentAI — Plano Gratuito")

    output = io.BytesIO()
    document.save(output)
    return ContentFile(output.getvalue())
