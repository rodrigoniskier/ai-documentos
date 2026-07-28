import io
import os
from unittest.mock import patch

from django.apps import apps
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document

from .document_forms import NewDocumentProjectForm
from .document_models import DocumentProject
from .document_services import create_openai_client
from .models import Plan, User
from .product_plans import apply_paid_plan_limits
from .services import ensure_plans, provision_free_account


def docx_upload(name="modelo.docx"):
    output = io.BytesIO()
    document = Document()
    document.add_heading("Modelo acadêmico", level=1)
    document.add_paragraph("{{ titulo }}")
    document.add_paragraph("{{ conteudo_gerado }}")
    document.save(output)
    return SimpleUploadedFile(
        name,
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


class DocumentModelRegistrationTests(TestCase):
    def test_new_models_are_registered_in_platform_app(self):
        self.assertEqual(apps.get_model("platform_app", "DocumentProject"), DocumentProject)
        self.assertIsNotNone(apps.get_model("platform_app", "DocumentTemplate"))
        self.assertIsNotNone(apps.get_model("platform_app", "ReferenceDocument"))


class NewDocumentFormTests(TestCase):
    def test_context_is_required_when_no_reference_is_attached(self):
        form = NewDocumentProjectForm(
            data={
                "document_type": "EMENTA",
                "title": "Ementa de Imunologia",
                "extra_fields_json": "{}",
            },
            files={"template_file": docx_upload()},
        )
        self.assertFalse(form.is_valid())
        self.assertIn("course_context", form.errors)
        self.assertIn("institution_context", form.errors)

    def test_context_allows_generation_without_references(self):
        form = NewDocumentProjectForm(
            data={
                "document_type": "PLANO_ENSINO",
                "title": "Plano de Ensino",
                "course_context": "Curso de Medicina",
                "institution_context": "Componente do primeiro período, turma presencial.",
                "extra_fields_json": '{"disciplina": "Imunologia"}',
            },
            files={"template_file": docx_upload()},
        )
        self.assertTrue(form.is_valid(), form.errors)
        self.assertEqual(form.cleaned_data["extra_fields_json"]["disciplina"], "Imunologia")


class OpenAIConfigurationTests(TestCase):
    @override_settings(OPENAI_API_KEY="test-key")
    @patch("platform_app.document_services.OpenAI")
    def test_client_ignores_implicit_organization_and_project(self, openai_class):
        os.environ["OPENAI_ORG_ID"] = "org-antiga"
        os.environ["OPENAI_PROJECT_ID"] = "proj-antigo"
        try:
            create_openai_client()
            openai_class.assert_called_once_with(
                api_key="test-key",
                timeout=120,
                max_retries=2,
            )
            self.assertEqual(os.environ["OPENAI_ORG_ID"], "org-antiga")
            self.assertEqual(os.environ["OPENAI_PROJECT_ID"], "proj-antigo")
        finally:
            os.environ.pop("OPENAI_ORG_ID", None)
            os.environ.pop("OPENAI_PROJECT_ID", None)


class PaidPlanLimitTests(TestCase):
    def test_paid_limits_are_increased(self):
        ensure_plans()
        apply_paid_plan_limits()
        pro = Plan.objects.get(code="PRO")
        ultra = Plan.objects.get(code="PREMIUM")
        self.assertEqual(pro.monthly_credits, 60)
        self.assertEqual(pro.source_limit, 60)
        self.assertEqual(pro.daily_limit, 20)
        self.assertEqual(ultra.monthly_credits, 180)
        self.assertEqual(ultra.source_limit, 200)
        self.assertEqual(ultra.daily_limit, 60)


class DocumentWorkflowViewTests(TestCase):
    def setUp(self):
        ensure_plans()
        self.user = User.objects.create_user(
            email="workflow@example.com",
            password="SenhaForte123!",
            full_name="Professora Workflow",
            professional_name="Profa. Workflow",
        )
        provision_free_account(self.user)
        self.client.force_login(self.user)

    def test_workspace_is_the_main_dashboard(self):
        response = self.client.get(reverse("dashboard"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Criar novo documento")
        self.assertContains(response, "Prepare sua primeira geração")

    @patch("platform_app.document_views.generate_project_content")
    @patch("platform_app.document_views.process_template")
    def test_user_can_generate_editable_project(self, process_template, generate_content):
        generate_content.return_value = (
            {
                "title": "Plano de Ensino — Imunologia",
                "warnings": ["Revisar bibliografia."],
                "sections": [
                    {"heading": "Ementa", "body": "Fundamentos da resposta imune."},
                    {"heading": "Objetivos", "body": "Compreender mecanismos imunológicos."},
                ],
            },
            100,
            80,
        )
        response = self.client.post(
            reverse("project_new"),
            data={
                "document_type": "PLANO_ENSINO",
                "title": "Plano de Ensino",
                "course_context": "Medicina",
                "institution_context": "Primeiro período, ensino presencial.",
                "extra_fields_json": '{"disciplina": "Imunologia"}',
                "template_file": docx_upload(),
            },
        )
        project = DocumentProject.objects.get(owner=self.user)
        self.assertRedirects(response, reverse("project_edit", args=[project.pk]))
        self.assertEqual(project.status, "ready")
        self.assertEqual(project.content["sections"][0]["heading"], "Ementa")
        self.user.wallet.refresh_from_db()
        self.assertEqual(self.user.wallet.balance, 3)
        process_template.assert_called_once()

        project.content["sections"][0]["body"] = "Ensino & pesquisa <integrados>"
        project.save(update_fields=["content"])
        pdf_response = self.client.get(reverse("project_download_pdf", args=[project.pk]))
        self.assertEqual(pdf_response.status_code, 200)
        pdf_bytes = b"".join(pdf_response.streaming_content)
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))
