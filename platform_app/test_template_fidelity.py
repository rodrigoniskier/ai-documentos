import io
from pathlib import Path

from django.conf import settings
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase
from django.urls import reverse
from docx import Document
from docx.shared import Cm

from .document_export import build_project_docx
from .document_models import DocumentProject, DocumentTemplate
from .models import User
from .owner_account import OWNER_CAPACITY, ensure_owner_account
from .services import ensure_plans, provision_free_account


def formatted_template_upload():
    output = io.BytesIO()
    document = Document()
    section = document.sections[0]
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.9)

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.add_run("PLANO DE ENSINO").bold = True

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Disciplina"
    table.cell(0, 1).text = "Disciplina de exemplo"

    document.add_heading("1. Ementa", level=2)
    document.add_paragraph("Texto antigo que deverá ser substituído.")
    document.sections[0].footer.paragraphs[0].text = "Rodapé institucional preservado"
    document.save(output)
    return SimpleUploadedFile(
        "modelo-formatado.docx",
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


class TemplateFidelityTests(TestCase):
    def setUp(self):
        ensure_plans()
        self.user = User.objects.create_user(
            email="fidelidade@example.com",
            password="SenhaForte123!",
            full_name="Professora Fidelidade",
            professional_name="Profa. Fidelidade",
        )
        provision_free_account(self.user)

    def test_docx_reuses_original_layout_and_formatting(self):
        template = DocumentTemplate.objects.create(
            owner=self.user,
            title="Modelo formatado",
            document_type="PLANO_ENSINO",
            file=formatted_template_upload(),
        )
        project = DocumentProject.objects.create(
            owner=self.user,
            template=template,
            document_type="PLANO_ENSINO",
            title="Plano de Ensino — Imunologia",
            course_context="Medicina",
            institution_context="Primeiro período",
            field_values={"disciplina": "Imunologia"},
            content={
                "title": "Plano de Ensino — Imunologia",
                "warnings": [],
                "sections": [
                    {
                        "heading": "Ementa",
                        "body": "Fundamentos da resposta imune inata e adaptativa.",
                    }
                ],
            },
            status="ready",
        )

        generated = build_project_docx(project, watermark=False)
        generated.seek(0)
        document = Document(io.BytesIO(generated.read()))

        self.assertAlmostEqual(document.sections[0].left_margin.cm, 1.7, places=1)
        self.assertAlmostEqual(document.sections[0].right_margin.cm, 1.9, places=1)
        self.assertEqual(document.tables[0].style.name, "Table Grid")
        self.assertEqual(document.tables[0].cell(0, 1).text, "Imunologia")
        self.assertIn(
            "Fundamentos da resposta imune inata e adaptativa.",
            "\n".join(paragraph.text for paragraph in document.paragraphs),
        )
        self.assertTrue(document.paragraphs[0].runs[0].bold)
        self.assertIn(
            "Rodapé institucional preservado",
            " ".join(
                paragraph.text
                for paragraph in document.sections[0].footer.paragraphs
            ),
        )


class OwnerAccountTests(TestCase):
    def setUp(self):
        ensure_plans()
        self.user = User.objects.create_user(
            email="niskier.rodrigo@gmail.com",
            password="SenhaForte123!",
            full_name="Rodrigo Niskier",
            professional_name="Prof. Rodrigo Niskier",
        )
        provision_free_account(self.user)

    def test_owner_account_is_promoted_to_internal_unlimited_plan(self):
        self.assertTrue(ensure_owner_account(self.user))
        self.user.subscription.refresh_from_db()
        self.user.wallet.refresh_from_db()
        self.assertEqual(self.user.subscription.plan.code, "OWNER")
        self.assertEqual(self.user.subscription.plan.daily_limit, OWNER_CAPACITY)
        self.assertEqual(self.user.subscription.plan.source_limit, OWNER_CAPACITY)
        self.assertFalse(self.user.subscription.plan.watermark)
        self.assertEqual(self.user.wallet.balance, OWNER_CAPACITY)

    def test_dashboard_displays_unlimited_credits(self):
        self.client.force_login(self.user)
        response = self.client.get(reverse("dashboard"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Conta proprietária sem limite de gerações")
        self.assertContains(response, "∞")


class PublicEmailTests(TestCase):
    def test_public_pages_use_new_support_email(self):
        ensure_plans()
        response = self.client.get(reverse("home"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "backuprodrigoniskier@gmail.com")
        self.assertNotContains(response, "rncontentlab@gmail.com")

    def test_old_email_is_absent_from_project_text_files(self):
        ignored_parts = {".git", "__pycache__", "staticfiles"}
        extensions = {".py", ".html", ".md", ".txt", ".example", ".yml", ".yaml"}
        occurrences = []
        for path in Path(settings.BASE_DIR).rglob("*"):
            if not path.is_file() or any(part in ignored_parts for part in path.parts):
                continue
            if path.suffix.lower() not in extensions and path.name != ".env.example":
                continue
            try:
                content = path.read_text(encoding="utf-8")
            except (UnicodeDecodeError, OSError):
                continue
            if "rncontentlab@gmail.com" in content.lower():
                occurrences.append(str(path.relative_to(settings.BASE_DIR)))
        self.assertEqual(occurrences, [])
