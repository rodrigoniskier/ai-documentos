import io
import re
import unicodedata

from django.conf import settings
from django.core.files.base import ContentFile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .document_conversion import convert_docx_bytes_to_pdf, libreoffice_available
from .document_quality import audit_docx_structure, inspect_visual_fidelity
from .template_exact_fields import fill_exact_adjacent_fields
from .template_renderer import render_project_in_template, replace_paragraph_text


class DocumentFidelityError(RuntimeError):
    pass


def _slug(value):
    normalized = unicodedata.normalize("NFKD", str(value or ""))
    normalized = normalized.encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^a-z0-9]+", "_", normalized.lower()).strip("_")


def _template_values(project):
    values = {
        str(key): value
        for key, value in (project.field_values or {}).items()
        if value not in (None, "")
    }
    values.update(
        {
            "titulo": project.content.get("title") or project.title,
            "curso": project.course_context,
            "contexto": project.institution_context,
            "tipo_documento": project.get_document_type_display(),
        }
    )
    for item in (project.content or {}).get("resolved_fields", []):
        if isinstance(item, dict) and item.get("value") not in (None, ""):
            key = item.get("key") or item.get("label")
            if key:
                values[str(key)] = item["value"]
    full_text = []
    for section in project.content.get("sections", []):
        heading = str(section.get("heading") or "Seção")
        body = str(section.get("body") or "")
        values[_slug(heading)] = body
        full_text.append(f"{heading}\n{body}")
    values["conteudo_gerado"] = "\n\n".join(full_text)
    return values


def _replace_placeholders_in_paragraph(paragraph, values):
    original = paragraph.text
    updated = original
    for key, value in values.items():
        updated = re.sub(
            r"\{\{\s*" + re.escape(str(key)) + r"\s*\}\}",
            str(value),
            updated,
            flags=re.I,
        )
    if updated == original:
        return False
    return replace_paragraph_text(paragraph, updated)


def _replace_placeholders(document, project):
    values = _template_values(project)
    changes = 0

    def replace_in_container(paragraphs, tables):
        nonlocal changes
        for paragraph in paragraphs:
            changes += int(_replace_placeholders_in_paragraph(paragraph, values))
        for table in tables:
            seen = set()
            for row in table.rows:
                for cell in row.cells:
                    identity = id(cell._tc)
                    if identity in seen:
                        continue
                    seen.add(identity)
                    replace_in_container(cell.paragraphs, cell.tables)

    replace_in_container(document.paragraphs, document.tables)
    for section in document.sections:
        replace_in_container(section.header.paragraphs, section.header.tables)
        replace_in_container(section.footer.paragraphs, section.footer.tables)
    return changes


def _append_generated_content(document, project):
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    if project.logo:
        try:
            project.logo.open("rb")
            document.add_picture(project.logo, width=Cm(3))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        finally:
            project.logo.close()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(project.content.get("title") or project.title).bold = True
    for warning in project.content.get("warnings", []):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"Atenção: {warning}").italic = True
    for section in project.content.get("sections", []):
        document.add_heading(section.get("heading") or "Seção", level=2)
        body = str(section.get("body") or "—")
        for line in body.splitlines() or ["—"]:
            document.add_paragraph(line or " ")


def _add_free_watermark(document):
    for section in document.sections:
        footer = section.footer
        existing = " ".join(paragraph.text for paragraph in footer.paragraphs)
        if "AjudAI Docente — Plano Gratuito" in existing:
            continue
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run("AjudAI Docente — Plano Gratuito").italic = True


def _serialize(document):
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _read_template_bytes(project):
    project.template.file.open("rb")
    try:
        return project.template.file.read()
    finally:
        project.template.file.close()


def _render_docx_bytes(project, watermark, *, compact=False):
    is_docx_template = project.template.file.name.lower().endswith(".docx")
    if is_docx_template:
        document, structural_changes = render_project_in_template(
            project, compact=compact
        )
        adjacent_changes = fill_exact_adjacent_fields(
            document, project, compact=compact
        )
        placeholder_changes = _replace_placeholders(document, project)
        if structural_changes + adjacent_changes + placeholder_changes == 0:
            raise DocumentFidelityError(
                "O modelo DOCX não possui marcadores nem blocos estruturais reconhecíveis. "
                "Para evitar alterar a diagramação de forma insegura, o documento não foi reconstruído."
            )
    else:
        document = Document()
        _append_generated_content(document, project)

    if watermark:
        _add_free_watermark(document)
    return _serialize(document)


def _verified_artifacts(project, watermark=False):
    is_docx_template = project.template.file.name.lower().endswith(".docx")
    strict_docx = _render_docx_bytes(project, watermark, compact=False)
    output_pdf = None
    report = None

    if not is_docx_template:
        if libreoffice_available():
            output_pdf = convert_docx_bytes_to_pdf(strict_docx)
        return strict_docx, output_pdf, report

    template_docx = _read_template_bytes(project)
    structure = audit_docx_structure(template_docx, strict_docx)
    if not structure.passed:
        raise DocumentFidelityError(
            "A auditoria estrutural impediu a exportação: " + " ".join(structure.issues)
        )

    if not libreoffice_available():
        return strict_docx, None, structure

    template_pdf = convert_docx_bytes_to_pdf(template_docx)
    strict_pdf = convert_docx_bytes_to_pdf(strict_docx)
    report = inspect_visual_fidelity(template_pdf, strict_pdf)
    selected_docx = strict_docx
    output_pdf = strict_pdf

    if not report.passed:
        compact_docx = _render_docx_bytes(project, watermark, compact=True)
        compact_structure = audit_docx_structure(template_docx, compact_docx)
        if compact_structure.passed:
            compact_pdf = convert_docx_bytes_to_pdf(compact_docx)
            compact_report = inspect_visual_fidelity(template_pdf, compact_pdf)
            if compact_report.passed or compact_report.score > report.score:
                selected_docx = compact_docx
                output_pdf = compact_pdf
                report = compact_report

    if (
        not report.passed
        and getattr(settings, "DOCUMENT_VISUAL_QA_BLOCK_ON_FAILURE", True)
    ):
        details = " ".join(report.issues[:5]) or report.summary
        raise DocumentFidelityError(
            "A inspeção visual final detectou perda de fidelidade e bloqueou o arquivo. "
            + details
        )
    return selected_docx, output_pdf, report


def build_project_docx(project, watermark=False):
    docx_bytes, _, _ = _verified_artifacts(project, watermark=watermark)
    return ContentFile(docx_bytes)


def build_project_pdf(project, watermark=False):
    docx_bytes, pdf_bytes, _ = _verified_artifacts(project, watermark=watermark)
    if pdf_bytes is None:
        pdf_bytes = convert_docx_bytes_to_pdf(docx_bytes)
    return ContentFile(pdf_bytes)
