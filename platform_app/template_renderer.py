import io
import re
import unicodedata
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


DOCUMENT_LABEL_ALIASES = {
    "titulo": ("titulo", "nome do documento"),
    "curso": ("curso", "area", "curso ou area"),
    "componente": (
        "componente",
        "disciplina",
        "componente curricular",
        "unidade curricular",
    ),
    "periodo": ("periodo", "etapa"),
    "turno": ("turno",),
    "modalidade": ("modalidade",),
    "semestre": ("semestre",),
    "professor responsavel": (
        "professor responsavel",
        "professor",
        "docente",
    ),
    "c h semanal": (
        "c h semanal",
        "ch semanal",
        "carga horaria semanal",
    ),
    "c h semestral": (
        "c h semestral",
        "ch semestral",
        "carga horaria semestral",
        "carga horaria total",
        "carga horaria",
    ),
    "contexto": ("contexto", "contexto institucional", "instituicao"),
    "tipo documento": ("tipo", "tipo de documento"),
    "objetivo geral": ("objetivo geral",),
    "objetivos": ("objetivos", "objetivos especificos"),
    "metodologia": ("metodologia", "estrategias metodologicas"),
    "avaliacao": ("avaliacao", "processo avaliativo", "estrategia avaliativa"),
    "bibliografia": ("bibliografia", "referencias", "referencias bibliograficas"),
    "observacoes": ("observacoes", "observacao"),
    "datas": ("datas", "data", "semanas"),
    "atividades": ("atividades", "atividades previstas"),
    "tema": ("tema", "assunto"),
    "duracao": ("duracao", "tempo previsto"),
    "recursos": ("recursos", "recursos didaticos"),
    "nivel": ("nivel", "nivel de dificuldade"),
    "quantidade itens": ("quantidade de itens", "numero de itens"),
    "tipo questoes": ("tipos de questoes", "tipo de questoes"),
    "criterios": ("criterios", "criterios de avaliacao"),
    "resultados": ("resultados",),
    "recomendacoes": ("recomendacoes",),
    "publico alvo": ("publico alvo",),
    "cronograma": ("cronograma",),
    "resultados esperados": ("resultados esperados",),
}

SECTION_LABEL_ALIASES = {
    "plano de ensino": ("plano de ensino", "identificacao", "dados gerais"),
    "ementa": ("ementa",),
    "objetivos competencias": (
        "objetivos competencias",
        "objetivos e competencias",
        "competencias",
        "objetivos",
    ),
    "conteudo": ("conteudo", "conteudo programatico"),
    "estrategias de ensino": (
        "estrategias de ensino",
        "metodologia",
        "estrategias metodologicas",
    ),
    "recursos disponiveis": (
        "recursos disponiveis",
        "recursos",
        "recursos didaticos",
    ),
    "avaliacao": ("avaliacao", "processo avaliativo"),
    "bibliografia basica": ("bibliografia basica",),
    "bibliografia complementar": ("bibliografia complementar",),
    "pontos para revisao docente": (
        "pontos para revisao docente",
        "pontos de revisao",
        "revisao docente",
        "observacoes",
    ),
}

_BULLET_PREFIX = re.compile(r"^\s*(?:[•▪◦●\-–—]\s*)+")
_UNIT_WITH_HOURS = re.compile(
    r"^\s*(UNIDADE\s+[IVXLCDM]+.*?)\s*(?:\|\s*)?"
    r"C\s*/?\s*H\s*:\s*([^\n|]+)\s*$",
    re.I,
)
_NUMBERING_PREFIX = re.compile(
    r"^\s*(?:(?:\d+)|(?:[IVXLCDM]+))\s*[.\-–—):]\s*",
    re.I,
)


def normalize_label(value):
    """Normaliza rótulos sem apagar palavras iniciadas por C, D, I, L, M, V ou X.

    A implementação anterior removia qualquer sequência inicial formada por algarismos
    romanos. Assim, palavras como ``CONTEÚDO`` e ``Cognitivos`` eram mutiladas e
    acabavam acionando correspondências incorretas.
    """

    value = unicodedata.normalize("NFKD", str(value or ""))
    value = value.encode("ascii", "ignore").decode("ascii")
    value = _NUMBERING_PREFIX.sub("", value)
    value = re.sub(r"[^a-zA-Z0-9]+", " ", value).strip().lower()
    return re.sub(r"\s+", " ", value)


def _read_docx(file_field):
    file_field.open("rb")
    try:
        payload = file_field.read()
    finally:
        file_field.close()
    return Document(io.BytesIO(payload))


def _common_prefix_length(left, right):
    limit = min(len(left), len(right))
    index = 0
    while index < limit and left[index] == right[index]:
        index += 1
    return index


def replace_paragraph_text(paragraph, value):
    """Substitui texto preservando o estilo do parágrafo e seus runs principais."""

    new_text = str(value or "")
    old_text = paragraph.text
    if old_text == new_text:
        return False

    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return True

    prefix_length = _common_prefix_length(old_text, new_text)
    if prefix_length == 0:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
        return True

    cursor = 0
    inserted = False
    for run in runs:
        original = run.text
        end = cursor + len(original)
        if end <= prefix_length:
            cursor = end
            continue
        if cursor < prefix_length:
            keep = prefix_length - cursor
            run.text = original[:keep] + new_text[prefix_length:]
            inserted = True
        elif not inserted:
            run.text = new_text[prefix_length:]
            inserted = True
        else:
            run.text = ""
        cursor = end

    if not inserted:
        runs[-1].text += new_text[prefix_length:]
    return True


def _append_paragraph_after(paragraph, text=""):
    new_element = deepcopy(paragraph._p)
    for text_node in new_element.xpath(".//w:t"):
        text_node.text = ""
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    replace_paragraph_text(new_paragraph, text)
    return new_paragraph


def _remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def _unique_cells(row):
    seen = set()
    result = []
    for cell in row.cells:
        identity = id(cell._tc)
        if identity in seen:
            continue
        seen.add(identity)
        result.append(cell)
    return result


def _clean_line(value):
    return _BULLET_PREFIX.sub("", str(value or "")).strip()


def _compact_paragraph(paragraph, floor=8.0, reduction=0.75):
    for run in paragraph.runs:
        if run.font.size is None:
            continue
        current = run.font.size.pt
        if current > floor:
            run.font.size = Pt(max(floor, current - reduction))


def _set_cell_lines(cell, lines, *, list_mode=True, join_plain=False, compact=False):
    cleaned = [_clean_line(line) for line in lines if _clean_line(line)]
    paragraphs = list(cell.paragraphs)

    if join_plain:
        text = " ".join(cleaned)
        target = next(
            (paragraph for paragraph in paragraphs if paragraph.text.strip()),
            paragraphs[0] if paragraphs else cell.add_paragraph(),
        )
        replace_paragraph_text(target, text)
        if compact:
            _compact_paragraph(target)
        target_element = target._p
        for paragraph in list(cell.paragraphs):
            if paragraph._p is not target_element and paragraph.text.strip():
                _remove_paragraph(paragraph)
        return 1

    content_paragraphs = [paragraph for paragraph in paragraphs if paragraph.text.strip()]
    template = (
        content_paragraphs[0]
        if content_paragraphs
        else paragraphs[-1] if paragraphs else cell.add_paragraph()
    )
    if list_mode:
        template = next(
            (
                paragraph
                for paragraph in paragraphs
                if "list" in (getattr(paragraph.style, "name", "") or "").lower()
            ),
            template,
        )

    targets = list(content_paragraphs) or [template]
    while len(targets) < len(cleaned):
        targets.append(_append_paragraph_after(targets[-1]))

    for index, paragraph in enumerate(list(targets)):
        if index < len(cleaned):
            replace_paragraph_text(paragraph, cleaned[index])
            if compact:
                _compact_paragraph(paragraph)
        else:
            _remove_paragraph(paragraph)

    if not cleaned:
        for paragraph in targets:
            if paragraph._p is not None:
                replace_paragraph_text(paragraph, "")
    return 1


def _canonical_field(label):
    normalized = normalize_label(label)
    if not normalized:
        return None
    for canonical, aliases in DOCUMENT_LABEL_ALIASES.items():
        normalized_aliases = {normalize_label(alias) for alias in aliases}
        if normalized in normalized_aliases:
            return normalize_label(canonical)
    return None


def _canonical_section(label):
    normalized = normalize_label(label)
    if not normalized:
        return None
    for canonical, aliases in SECTION_LABEL_ALIASES.items():
        normalized_aliases = {normalize_label(alias) for alias in aliases}
        if normalized in normalized_aliases:
            return normalize_label(canonical)
    return None


def _project_values(project):
    values = {
        _canonical_field(key) or normalize_label(key): str(value)
        for key, value in (project.field_values or {}).items()
        if value not in (None, "")
    }
    values.update(
        {
            "titulo": project.content.get("title") or project.title,
            "curso": project.course_context,
            "contexto": project.institution_context,
            "tipo documento": project.get_document_type_display(),
        }
    )
    for item in (project.content or {}).get("resolved_fields", []):
        if not isinstance(item, dict):
            continue
        key = _canonical_field(item.get("label") or item.get("key"))
        value = str(item.get("value") or "").strip()
        if key and value:
            values[key] = value
    return {key: value for key, value in values.items() if str(value).strip()}


def _section_values(project):
    result = {}
    for section in (project.content or {}).get("sections", []):
        heading = _canonical_section(section.get("heading"))
        body = str(section.get("body") or "").strip()
        if heading and body:
            result[heading] = body
    return result


def _parse_metadata(body):
    values = {}
    for line in str(body or "").splitlines():
        for part in re.split(r"\s*\|\s*", line):
            if ":" not in part:
                continue
            label, value = part.split(":", 1)
            canonical = _canonical_field(label)
            if canonical and value.strip():
                values[canonical] = value.strip()
    return values


def _metadata_values(project_values, section_values):
    """Prioriza valores resolvidos pela IA/modelo sobre campos brutos conflitantes."""

    metadata = dict(project_values)
    identification = section_values.get("plano de ensino")
    if identification:
        metadata.update(_parse_metadata(identification))
    return metadata


def _replace_labeled_paragraph(paragraph, value):
    text = paragraph.text
    if ":" not in text:
        return False
    colon_position = text.find(":") + 1
    cursor = 0
    inserted = False
    for run in paragraph.runs:
        original = run.text
        end = cursor + len(original)
        if end <= colon_position:
            cursor = end
            continue
        if cursor < colon_position:
            keep = colon_position - cursor
            run.text = original[:keep] + " " + str(value)
            inserted = True
        elif not inserted:
            run.text = " " + str(value)
            inserted = True
        else:
            run.text = ""
        cursor = end
    if not inserted:
        paragraph.add_run(" " + str(value))
    return True


def _fill_metadata_tables(document, metadata, *, compact=False):
    changes = 0
    processed = set()
    for table in document.tables:
        for row in table.rows:
            for cell in _unique_cells(row):
                identity = id(cell._tc)
                if identity in processed:
                    continue
                processed.add(identity)
                for paragraph in cell.paragraphs:
                    if ":" not in paragraph.text:
                        continue
                    label = paragraph.text.split(":", 1)[0]
                    canonical = _canonical_field(label)
                    if not canonical or canonical not in metadata:
                        continue
                    changes += int(
                        _replace_labeled_paragraph(paragraph, metadata[canonical])
                    )
                    if compact:
                        _compact_paragraph(paragraph)
    return changes


def _parse_subsections(body, labels):
    normalized_labels = {normalize_label(label) for label in labels}
    result = {label: [] for label in normalized_labels}
    current = None
    for raw_line in str(body or "").splitlines():
        line = raw_line.strip()
        normalized = normalize_label(line)
        if normalized in normalized_labels:
            current = normalized
            continue
        cleaned = _clean_line(line)
        if current and cleaned:
            result[current].append(cleaned)
    return result


def _parse_units(body):
    units = []
    current = None
    for raw_line in str(body or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = _UNIT_WITH_HOURS.match(line)
        if not match:
            alternative = re.match(
                r"^(.*?)\s*\|\s*C\s*/?\s*H\s*:\s*(.+)$",
                line,
                flags=re.I,
            )
            if alternative and normalize_label(alternative.group(1)).startswith("unidade"):
                match = alternative
        if match:
            current = {
                "name": match.group(1).strip(),
                "hours": match.group(2).strip(),
                "topics": [],
            }
            units.append(current)
            continue
        if current is not None:
            cleaned = _clean_line(line)
            if cleaned:
                current["topics"].append(cleaned)
    return units


def _table_heading(table):
    if not table.rows:
        return None
    for cell in _unique_cells(table.rows[0]):
        for paragraph in cell.paragraphs:
            if not paragraph.text.strip():
                continue
            heading = _canonical_section(paragraph.text)
            if heading:
                return heading
            break
    return None


def _fill_objectives_table(table, body, *, compact=False):
    subsections = _parse_subsections(body, ("Cognitivos", "Habilidades", "Atitudes"))
    changes = 0
    for row in table.rows[1:]:
        cells = _unique_cells(row)
        if len(cells) < 2:
            continue
        label = normalize_label(cells[0].text)
        if label in subsections and subsections[label]:
            changes += _set_cell_lines(
                cells[1], subsections[label], list_mode=True, compact=compact
            )
    return changes


def _fill_content_table(table, body, *, compact=False):
    units = _parse_units(body)
    if not units or len(table.rows) < 3:
        return 0
    changes = 0
    for index, row in enumerate(table.rows[2:]):
        cells = _unique_cells(row)
        if len(cells) < 3:
            continue
        if index >= len(units):
            # Não elimina linhas estruturais quando a IA devolve menos unidades.
            continue
        unit = units[index]
        changes += _set_cell_lines(
            cells[0], [unit["name"]], list_mode=False, join_plain=True, compact=compact
        )
        changes += _set_cell_lines(
            cells[1], [unit["hours"]], list_mode=False, join_plain=True, compact=compact
        )
        changes += _set_cell_lines(
            cells[2], unit["topics"], list_mode=True, compact=compact
        )
    return changes


def _fill_bibliography_table(table, section_values, *, compact=False):
    if len(table.rows) < 2:
        return 0
    headers = _unique_cells(table.rows[0])
    targets = _unique_cells(table.rows[1])
    changes = 0
    for index, header in enumerate(headers):
        if index >= len(targets):
            continue
        heading = None
        for paragraph in header.paragraphs:
            heading = _canonical_section(paragraph.text)
            if heading:
                break
        if heading and section_values.get(heading):
            changes += _set_cell_lines(
                targets[index],
                section_values[heading].splitlines(),
                list_mode=True,
                compact=compact,
            )
    return changes


def _fill_review_table(table, body, *, compact=False):
    if not table.rows:
        return 0
    cell = _unique_cells(table.rows[0])[0]
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        return 0
    lines = [_clean_line(line) for line in body.splitlines() if _clean_line(line)]
    targets = list(paragraphs[1:])
    if not targets:
        targets = [_append_paragraph_after(paragraphs[0])]
    while len(targets) < len(lines):
        targets.append(_append_paragraph_after(targets[-1]))
    for index, paragraph in enumerate(list(targets)):
        if index < len(lines):
            replace_paragraph_text(paragraph, lines[index])
            if compact:
                _compact_paragraph(paragraph)
        else:
            _remove_paragraph(paragraph)
    return 1


def _fill_section_tables(document, section_values, *, compact=False):
    changes = 0
    for table in document.tables:
        heading = _table_heading(table)
        if not heading:
            continue
        body = section_values.get(heading)

        if heading == "bibliografia basica":
            changes += _fill_bibliography_table(
                table, section_values, compact=compact
            )
            continue
        if not body:
            continue
        if heading == "ementa" and len(table.rows) >= 2:
            target = _unique_cells(table.rows[1])[0]
            changes += _set_cell_lines(
                target,
                [body],
                list_mode=False,
                join_plain=True,
                compact=compact,
            )
        elif heading == "objetivos competencias":
            changes += _fill_objectives_table(table, body, compact=compact)
        elif heading == "conteudo":
            changes += _fill_content_table(table, body, compact=compact)
        elif heading in {
            "estrategias de ensino",
            "recursos disponiveis",
            "avaliacao",
        } and len(table.rows) >= 2:
            target = _unique_cells(table.rows[1])[0]
            changes += _set_cell_lines(
                target, body.splitlines(), list_mode=True, compact=compact
            )
        elif heading == "pontos para revisao docente":
            changes += _fill_review_table(table, body, compact=compact)
    return changes


def _fill_paragraph_sections(document, section_values, *, compact=False):
    """Preenche modelos sem tabelas usando somente títulos de seção exatos."""

    changes = 0
    paragraphs = list(document.paragraphs)
    headings = []
    for index, paragraph in enumerate(paragraphs):
        heading = _canonical_section(paragraph.text)
        if heading and heading != "plano de ensino" and heading in section_values:
            headings.append((index, heading))

    for position, (heading_index, heading) in enumerate(headings):
        end = headings[position + 1][0] if position + 1 < len(headings) else len(paragraphs)
        candidates = [
            paragraph
            for paragraph in paragraphs[heading_index + 1 : end]
            if paragraph.text.strip()
        ]
        lines = [
            _clean_line(line)
            for line in section_values[heading].splitlines()
            if _clean_line(line)
        ]
        if not candidates:
            paragraph = _append_paragraph_after(paragraphs[heading_index])
            candidates = [paragraph]
        while len(candidates) < len(lines):
            candidates.append(_append_paragraph_after(candidates[-1]))
        for index, paragraph in enumerate(list(candidates)):
            if index < len(lines):
                replace_paragraph_text(paragraph, lines[index])
                if compact:
                    _compact_paragraph(paragraph)
            else:
                _remove_paragraph(paragraph)
        changes += 1
    return changes


def inspect_template_layout(file_field):
    """Retorna um mapa compacto de layout para orientar a geração da IA."""

    if not getattr(file_field, "name", "").lower().endswith(".docx"):
        return {"format": "reference-only", "tables": [], "metadata_fields": []}

    document = _read_docx(file_field)
    tables = []
    metadata_fields = []
    for table_index, table in enumerate(document.tables):
        heading = _table_heading(table)
        slots = []
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(_unique_cells(row)):
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if ":" in text:
                        label = text.split(":", 1)[0]
                        canonical = _canonical_field(label)
                        if canonical:
                            metadata_fields.append(canonical)
                if row_index == 0:
                    continue
                text = cell.text.strip()
                if text:
                    slots.append(
                        {
                            "row": row_index,
                            "column": cell_index,
                            "capacity_chars": len(text),
                            "paragraphs": max(1, len([p for p in cell.paragraphs if p.text.strip()])),
                        }
                    )
        tables.append(
            {
                "index": table_index,
                "heading": heading or "",
                "rows": len(table.rows),
                "columns": len(_unique_cells(table.rows[0])) if table.rows else 0,
                "slots": slots[:16],
            }
        )
    return {
        "format": "docx",
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "metadata_fields": sorted(set(metadata_fields)),
        "tables": tables,
        "instructions": (
            "Mantenha exatamente os títulos das seções, a quantidade de unidades e "
            "a função de cada coluna. Produza texto compatível com a capacidade aproximada "
            "dos blocos para evitar expansão desnecessária da paginação."
        ),
    }


def render_project_in_template(project, *, compact=False):
    """Preenche o DOCX original sem reconstruir ou adivinhar sua estrutura visual."""

    document = _read_docx(project.template.file)
    project_values = _project_values(project)
    section_values = _section_values(project)
    metadata = _metadata_values(project_values, section_values)

    changes = 0
    changes += _fill_metadata_tables(document, metadata, compact=compact)
    changes += _fill_section_tables(document, section_values, compact=compact)
    changes += _fill_paragraph_sections(document, section_values, compact=compact)
    return document, changes


def clone_paragraph_format(source, target):
    """Utilitário compatível com testes e futuras expansões de blocos dinâmicos."""

    target._p.get_or_add_pPr().clear_content()
    if source._p.pPr is not None:
        target._p.insert(0, deepcopy(source._p.pPr))
