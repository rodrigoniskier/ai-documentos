import io
import json
import os
import re
import unicodedata
from contextlib import contextmanager
from html import escape

from django.conf import settings
from django.core.files.base import ContentFile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openai import AuthenticationError, OpenAI
from pypdf import PdfReader
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


DOCUMENT_SCHEMA = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "warnings": {"type": "array", "items": {"type": "string"}},
        "sections": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "heading": {"type": "string"},
                    "body": {"type": "string"},
                },
                "required": ["heading", "body"],
                "additionalProperties": False,
            },
        },
    },
    "required": ["title", "warnings", "sections"],
    "additionalProperties": False,
}


@contextmanager
def _without_implicit_openai_scope():
    """Impede que variáveis antigas vinculem a requisição a outra organização."""
    removed = {}
    for name in ("OPENAI_ORG_ID", "OPENAI_PROJECT_ID"):
        if name in os.environ:
            removed[name] = os.environ.pop(name)
    try:
        yield
    finally:
        os.environ.update(removed)


def create_openai_client():
    if not settings.OPENAI_API_KEY:
        raise RuntimeError("A chave da OpenAI ainda não foi configurada no servidor.")
    with _without_implicit_openai_scope():
        return OpenAI(api_key=settings.OPENAI_API_KEY, timeout=120, max_retries=2)


def _read_docx(file_obj):
    document = Document(file_obj)
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                lines.append(" | ".join(values))
    return "\n".join(lines)


def extract_uploaded_text(file_field, max_chars=100_000):
    file_field.open("rb")
    try:
        name = file_field.name.lower()
        if name.endswith(".pdf"):
            text = "\n".join((page.extract_text() or "") for page in PdfReader(file_field).pages)
        elif name.endswith(".docx"):
            text = _read_docx(file_field)
        else:
            text = file_field.read().decode("utf-8", errors="ignore")
        return text[:max_chars]
    finally:
        file_field.close()


def extract_placeholders(text):
    return sorted(
        {
            item.strip()
            for item in re.findall(r"\{\{\s*([^{}]+?)\s*\}\}", text)
            if item.strip()
        }
    )


def process_template(template):
    template.extracted_text = extract_uploaded_text(template.file)
    template.placeholders = extract_placeholders(template.extracted_text)
    template.save(update_fields=["extracted_text", "placeholders"])


def process_reference(reference):
    try:
        reference.extracted_text = extract_uploaded_text(reference.file)
        reference.status = "done"
    except Exception:
        reference.extracted_text = ""
        reference.status = "failed"
    reference.save(update_fields=["extracted_text", "status"])


def _truncate_sources(references, total_limit=120_000):
    result = []
    remaining = total_limit
    for reference in references:
        if remaining <= 0:
            break
        text = (reference.extracted_text or "")[:remaining]
        result.append({"title": reference.title, "text": text})
        remaining -= len(text)
    return result


def generate_project_content(project):
    references = list(project.references.filter(status="done"))
    no_reference_warning = not references
    payload = {
        "document_type": project.get_document_type_display(),
        "requested_title": project.title,
        "template_structure": project.template.extracted_text[:60_000],
        "template_placeholders": project.template.placeholders,
        "optional_fields": project.field_values,
        "course_context": project.course_context,
        "institution_context": project.institution_context,
        "references": _truncate_sources(references),
        "reference_limitation": (
            "Nenhuma referência foi anexada. Use apenas o contexto informado, declare a limitação e recomende revisão humana rigorosa."
            if no_reference_warning
            else "Use as referências como fonte prioritária e não invente dados ausentes."
        ),
    }
    system_prompt = (
        "Você personaliza documentos acadêmicos a partir de um modelo enviado pelo usuário. "
        "Respeite a ordem, os títulos e a finalidade percebida no modelo. Preencha apenas com "
        "informações sustentadas pelos campos ou referências fornecidas. Não invente normas, "
        "datas, bibliografias ou dados institucionais. Produza seções claras e editáveis. "
        "Quando faltarem referências, inclua um aviso explícito sobre a limitação e a revisão humana."
    )
    try:
        client = create_openai_client()
        response = client.responses.create(
            model=settings.OPENAI_MODEL,
            store=False,
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
            ],
            text={
                "format": {
                    "type": "json_schema",
                    "name": "documento_academico_editavel",
                    "strict": True,
                    "schema": DOCUMENT_SCHEMA,
                }
            },
        )
    except AuthenticationError as exc:
        code = getattr(exc, "code", "") or ""
        if code == "invalid_organization" or "organization" in str(exc).lower():
            raise RuntimeError(
                "A credencial da OpenAI está vinculada a uma organização ou projeto sem acesso. "
                "Remova OPENAI_ORG_ID e OPENAI_PROJECT_ID do Render ou substitua OPENAI_API_KEY por uma chave ativa do projeto correto."
            ) from exc
        raise
    usage = getattr(response, "usage", None)
    return (
        json.loads(response.output_text),
        getattr(usage, "input_tokens", 0) or 0,
        getattr(usage, "output_tokens", 0) or 0,
    )


def _slug(value):
    normalized = (
        unicodedata.normalize("NFKD", value)
        .encode("ascii", "ignore")
        .decode("ascii")
    )
    return re.sub(r"[^a-z0-9]+", "_", normalized.lower()).strip("_")


def _replace_in_paragraph(paragraph, values):
    original = paragraph.text
    updated = original
    for key, value in values.items():
        updated = re.sub(
            r"\{\{\s*" + re.escape(key) + r"\s*\}\}",
            str(value),
            updated,
            flags=re.I,
        )
    if updated != original:
        paragraph.clear()
        paragraph.add_run(updated)
        return True
    return False


def _template_values(project):
    values = {
        str(key): value
        for key, value in project.field_values.items()
        if value not in (None, "")
    }
    values.update(
        {
            "titulo": project.content.get("title") or project.title,
            "curso": project.course_context,
            "contexto": project.institution_context,
            "tipo_documento": project.get_document_type_display(),
        }
    )
    full_text = []
    for section in project.content.get("sections", []):
        heading = str(section.get("heading") or "Seção")
        body = str(section.get("body") or "")
        values[_slug(heading)] = body
        full_text.append(f"{heading}\n{body}")
    values["conteudo_gerado"] = "\n\n".join(full_text)
    return values


def _base_docx(project):
    if project.template.file.name.lower().endswith(".docx"):
        project.template.file.open("rb")
        try:
            return Document(project.template.file)
        finally:
            project.template.file.close()
    return Document()


def build_project_docx(project, watermark=False):
    document = _base_docx(project)
    values = _template_values(project)
    replaced = False
    for paragraph in document.paragraphs:
        replaced = _replace_in_paragraph(paragraph, values) or replaced
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replaced = _replace_in_paragraph(paragraph, values) or replaced

    if not replaced:
        body = document._element.body
        for child in list(body):
            if not child.tag.endswith("sectPr"):
                body.remove(child)
        styles = document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(11)
        if project.logo:
            try:
                project.logo.open("rb")
                document.add_picture(project.logo, width=Cm(3))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            finally:
                project.logo.close()
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(project.content.get("title") or project.title).bold = True
        for warning in project.content.get("warnings", []):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"Atenção: {warning}")
            run.italic = True
        for section in project.content.get("sections", []):
            document.add_heading(section.get("heading") or "Seção", level=2)
            body_text = str(section.get("body") or "—")
            for line in body_text.splitlines() or ["—"]:
                document.add_paragraph(line or " ")

    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Documento revisável gerado no AjudAI Docente"
    if watermark:
        footer.add_run(" • Plano Gratuito")
    output = io.BytesIO()
    document.save(output)
    return ContentFile(output.getvalue())


def _pdf_markup(value):
    return escape(str(value), quote=False).replace("\n", "<br/>")


def build_project_pdf(project, watermark=False):
    output = io.BytesIO()
    story = []
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ProjectTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    if project.logo:
        try:
            project.logo.open("rb")
            logo_data = io.BytesIO(project.logo.read())
        finally:
            project.logo.close()
        story.append(
            PdfImage(logo_data, width=3 * cm, height=3 * cm, kind="proportional")
        )
        story.append(Spacer(1, 10))
    story.append(
        Paragraph(
            _pdf_markup(project.content.get("title") or project.title),
            title_style,
        )
    )
    for warning in project.content.get("warnings", []):
        story.append(
            Paragraph(
                f"<i>Atenção: {_pdf_markup(warning)}</i>",
                styles["BodyText"],
            )
        )
        story.append(Spacer(1, 8))
    for section in project.content.get("sections", []):
        story.append(
            Paragraph(
                _pdf_markup(section.get("heading") or "Seção"),
                styles["Heading2"],
            )
        )
        story.append(
            Paragraph(
                _pdf_markup(section.get("body") or "—"),
                styles["BodyText"],
            )
        )
        story.append(Spacer(1, 12))
    footer = "AjudAI Docente"
    if watermark:
        footer += " — Plano Gratuito"

    def add_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(A4[0] / 2, 1.2 * cm, footer)
        canvas.restoreState()

    pdf = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    pdf.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
    return ContentFile(output.getvalue())
